import sys
import streamlit as st
import os
import tempfile
from datetime import datetime
import pandas as pd
import docx
from docx import Document
import comtypes.client
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import fitz  # PyMuPDF
import io
import win32gui
import win32ui
import win32con
import win32com.client
from dotenv import load_dotenv
import pytesseract
import mammoth
import base64
import win32com.client
import pythoncom
import time
from PIL import Image
import openai
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl import Workbook, load_workbook
import re
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from docxcompose.composer import Composer
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# from trial import _extract_from_excel as EXTRACT_FROM_EXCEL
                

from typing import List, Dict, Tuple, Union, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)



load_dotenv(".env")

os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')

# # Set page configuration
# st.set_page_config(
#     page_title="AI Audit Report Generator",
#     page_icon="📊",
#     layout="wide",
#     initial_sidebar_state="expanded"
# )



# ------- Template Analysis Module -------
class TemplateAnalyzer:
    """Analyze audit report templates and extract their structure for direct filling."""
    
    @staticmethod
    def extract_template_structure(template_path):
        """Extract the detailed structure of the template document including form fields and tables."""
        try:
            doc = Document(template_path)
            template_structure = {
                'paragraphs': [],
                'tables': []
            }
            
            # Extract paragraph headers and any form fields
            for para in doc.paragraphs:
                if para.text.strip():
                    template_structure['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name,
                        'alignment': para.alignment
                    })
            
            # Extract tables with detailed structure
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'id': f'table_{table_idx}',
                    'rows': len(table.rows),
                    'cols': len(table.rows[0].cells) if table.rows else 0,
                    'cells': []
                }
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        table_data['cells'].append({
                            'row': row_idx,
                            'col': col_idx,
                            'text': cell_text,
                            'is_header': row_idx == 0 or col_idx == 0,
                            'is_form_field': cell_text == '' and (row_idx > 0 and col_idx > 0)
                        })
                
                template_structure['tables'].append(table_data)
            
            return template_structure
        except Exception as e:
            st.error(f"Error analyzing template: {e}")
            return {'paragraphs': [], 'tables': []}
    
    @staticmethod
    def format_template_for_prompt(template_structure):
        """Format the template structure for use in the LLM prompt."""
        prompt_parts = ["# Template Structure"]
        
        # Add paragraphs info
        prompt_parts.append("## Paragraphs and Headers:")
        for para in template_structure['paragraphs']:
            if para['style'].startswith('Heading'):
                level = para['style'].replace('Heading', '')
                level = int(level) if level.isdigit() else 1
                prompt_parts.append(f"{'#' * level} {para['text']}")
            else:
                prompt_parts.append(f"- {para['text']}")
        
        # Add tables info
        prompt_parts.append("## Tables:")
        for table in template_structure['tables']:
            prompt_parts.append(f"Table with {table['rows']} rows and {table['cols']} columns:")
            
            # Group cells by row to recreate table structure
            rows = {}
            for cell in table['cells']:
                if cell['row'] not in rows:
                    rows[cell['row']] = []
                rows[cell['row']].append(cell)
            
            # Format as markdown table
            for row_idx in sorted(rows.keys()):
                row_cells = sorted(rows[row_idx], key=lambda x: x['col'])
                row_text = "| " + " | ".join(cell['text'] if cell['text'] else '[EMPTY]' for cell in row_cells) + " |"
                prompt_parts.append(row_text)
                
                # Add separator after header row
                if row_idx == 0:
                    prompt_parts.append("| " + " | ".join(['---'] * len(row_cells)) + " |")
        
        return '\n'.join(prompt_parts)

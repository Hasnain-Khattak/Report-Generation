import sys
import streamlit as st
import os
import tempfile
from datetime import datetime
import pandas as pd
import docx
from docx import Document
import comtypes.client
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import fitz  # PyMuPDF
import io
import win32gui
import win32ui
import win32con
import win32com.client
from dotenv import load_dotenv
import pytesseract
import mammoth
import base64
import win32com.client
import pythoncom
import time
from PIL import Image
import openai
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl import Workbook, load_workbook
import re
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from docxcompose.composer import Composer
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# from trial import _extract_from_excel as EXTRACT_FROM_EXCEL
                

from typing import List, Dict, Tuple, Union, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)



load_dotenv(".env")

os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')

# Set page configuration
# st.set_page_config(
#     page_title="AI Audit Report Generator",
#     page_icon="📊",
#     layout="wide",
#     initial_sidebar_state="expanded"
# )

# ------- Document Processing Module -------
class DocumentProcessor:
    """Process various document types and extract text."""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """Extract text from PDF files."""
        try:
            text = ""
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            return text
        except Exception as e:
            st.error(f"Error extracting text from PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    full_text.append(" | ".join(row_text))
            
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_image(file_path):
        """Extract text from image files using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            st.error(f"Error extracting text from image: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_excel(file_path):
        """Extract text from Excel files."""
        try:
            # Use xlrd for .xls files
            if file_path.lower().endswith('.xls'):
                import xlrd
                workbook = xlrd.open_workbook(file_path)
                full_text = []
                
                for sheet_name in workbook.sheet_names():
                    sheet = workbook.sheet_by_name(sheet_name)
                    full_text.append(f"Sheet: {sheet_name}")
                    
                    # Process headers
                    headers = []
                    for col in range(sheet.ncols):
                        headers.append(str(sheet.cell_value(0, col)))
                    
                    # Process rows
                    for row in range(1, sheet.nrows):
                        row_data = []
                        for col in range(sheet.ncols):
                            row_data.append(str(sheet.cell_value(row, col)))
                        full_text.append(" | ".join(row_data))
                
                return '\n'.join(full_text)
            
            # Use openpyxl for .xlsx and .xlsm files
            else:
                import pandas as pd
                # Make sure pandas is installed and available
                try:
                    # Try to import openpyxl directly to ensure it's available
                    import openpyxl
                except ImportError:
                    st.error("Missing openpyxl library. Install it with: pip install openpyxl")
                    return "ERROR: openpyxl library missing. Install with: pip install openpyxl"
                
                df_dict = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
                
                full_text = []
                for sheet_name, df in df_dict.items():
                    full_text.append(f"Sheet: {sheet_name}")
                    full_text.append(df.to_string(index=False))
                
                return '\n'.join(full_text)
                
        except ImportError as e:
            st.error(f"Missing required library: {e}. Install openpyxl with: pip install openpyxl")
            return f"ERROR: Missing required library: {e}. Install with pip."
        except Exception as e:
            st.error(f"Error extracting text from Excel: {e}")
            return f"Error: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file_path):
        """Extract text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error extracting text from TXT: {e}")
            return ""
    
    @staticmethod
    def process_file(file_path):
        """Process a file based on its extension."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.pdf']:
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc', '.dotx', '.dot']:
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            return DocumentProcessor.extract_text_from_image(file_path)
        elif file_extension in ['.xlsx', '.xls', '.xlsm']:
            return DocumentProcessor.extract_text_from_excel(file_path)
        elif file_extension in ['.txt']:
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            return f"Unsupported file format: {file_extension}"

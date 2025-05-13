import os
import tempfile
import re
import time
import logging
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from dotenv import load_dotenv
import pythoncom

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

load_dotenv(".env")
os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')


class ReportGenerator:
    """Generate audit reports with evidence images and score-based checkmark placement."""
    
    def __init__(self):
        """Initialize the ReportGenerator."""
        self.evidence_images = {}
        self.evidence_metadata = {}
        self.preprocessor = None
        self.image_paths = []
        self.report_images_dir = None
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        self.process_list = []
        self.used_evidence = set()
        
    def set_evidence_images(self, evidence_images, evidence_metadata=None):
        """
        Set the evidence images and metadata to be used in the report.
        
        Args:
            evidence_images: Dictionary of evidence files and their images
            evidence_metadata: Dictionary with score and category information
        """
        self.evidence_images = evidence_images
        self.evidence_metadata = evidence_metadata or {}
    
    def set_preprocessor(self, preprocessor):
        """Set the response preprocessor instance."""
        self.preprocessor = preprocessor

    def clean_audit_data(self, raw_text: str):
        import re
        import logging
        logger = logging.getLogger(__name__)
        
        lines = raw_text.strip("`").splitlines()
        logger.info(f"UNCLEANED:{lines}")

        result = [
            {"header": {}, "legend": {}},
            {"body": []},
            {"footer": {}}
        ]

        section = "header"
        legend_keys = {"OK", "OFI", "NC", "NA"}
        process_table_active = False
        process_headers = []
        in_final_comments = False
        final_comments_text = []

        for line in lines:
            line = line.strip()
            if not line or line.startswith("```"):
                continue
                
            # Handle section headers and separators
            if line.startswith("# ") or line == "|---|---|":
                continue
                
            # Special handling for AUDIT REPORT FINAL COMMENTS section
            if line == "## AUDIT REPORT FINAL COMMENTS":
                in_final_comments = True
                section = "footer"
                continue
                
            if in_final_comments:
                # Check if we've reached the signature block at the end
                if "Internal Auditor" in line:
                    if final_comments_text:  # Only add if we have content
                        result[2]["footer"]["AUDIT REPORT FINAL COMMENTS"] = "\n".join(final_comments_text)
                    in_final_comments = False
                    continue
                
                # Collect final comments text
                if line:  # Only add non-empty lines
                    final_comments_text.append(line)
                continue

            parts = [p.strip() for p in line.strip("|").split("|")]

            # Handle Legend
            if section == "header" and len(parts) == 2 and parts[0] in legend_keys:
                result[0]["legend"][parts[0]] = parts[1]
                section = "legend"
                continue

            # Handle Process Table Header
            if len(parts) == 7 and parts[0].upper() == "PROCESS":
                process_headers = parts
                process_table_active = True
                section = "body"
                continue

            # Handle Process Table Rows
            if process_table_active and len(parts) == 7:
                process_data = {}
                for idx, key in enumerate(process_headers):
                    val = parts[idx].strip()
                    if key.upper().startswith("PROCESS"):
                        val = re.sub(r"^Process:\s*", "", val, flags=re.I)
                        key = "PROCESS"
                    elif key.upper().startswith("SIGHTED"):
                        val = re.sub(r"^Evidence:\s*", "", val, flags=re.I)
                        key = "SIGHTED EVIDENCE"
                    process_data[key.strip()] = val
                result[1]["body"].append(process_data)
                continue

            # Handle Key:Value pairs
            if len(parts) == 2:
                key, val = parts
                if section == "header":
                    result[0]["header"][key] = val
                elif section in {"legend", "body"}:
                    section = "footer"
                    result[2]["footer"][key] = val
                elif section == "footer":
                    result[2]["footer"][key] = val

        # Remove empty PROCESS rows
        cleaned_body = []
        for row in result[1]["body"]:
            process_val = row.get("PROCESS", "")
            if not (process_val.strip() == "" or process_val == '[EMPTY]' or process_val == 'EMPTY' or process_val.replace("-", "") == "" or process_val.replace("|", "") == ""):
                cleaned_body.append(row)
        result[1]["body"] = cleaned_body

        # Format header values with proper newlines
        formatted = {}
        header = result[0]["header"]
        for i, (key, value) in enumerate(header.items()):
            if i == 0:
                # Keep the first item (title) untouched
                formatted[key] = value
            else:
                # Add a newline before each dash item, except if already on a new line
                value = re.sub(r'(<br>\s*)+$', '', value.strip())

                # Step 2: Replace <br> tags in the middle with a single newline
                value = re.sub(r'(<br>\s*)+', '\n', value)

                # Step 3: Add newline before dashes (excluding the first dash)
                value = re.sub(r'\s*-\s+', r'\n- ', value)

                formatted[key] = value

        result[0]['header'] = formatted
        
        # Format the footer values properly
        for key, value in result[2]['footer'].items():
            if '<br>' in value:
                # Add a newline before each dash item, except if already on a new line
                value = re.sub(r'(<br>\s*)+$', '', value.strip())

                # Replace <br> tags in the middle with a single newline
                value = re.sub(r'(<br>\s*)+', '\n', value)

                # Add newline before dashes (excluding the first dash)
                value = re.sub(r'\s*-\s+', r'\n- ', value)
                
                result[2]['footer'][key] = value
        
        # Ensure we always have the AUDIT REPORT FINAL COMMENTS in the footer
        if "AUDIT REPORT FINAL COMMENTS" not in result[2]["footer"] and final_comments_text:
            result[2]["footer"]["AUDIT REPORT FINAL COMMENTS"] = "\n".join(final_comments_text)
        
        # Check if audit final comments are missing
        if "AUDIT REPORT FINAL COMMENTS" not in result[2]["footer"]:
            logger.warning("AUDIT REPORT FINAL COMMENTS section is missing!")

        return result
    
    def fill_template_document(self, template_path, report_content, auditor_name=""):
        """Fill the template document with the report content including evidence images."""
        try:
            # Load the template document
            doc = Document(template_path)
            
            # Clean the data
            # logger.info(f"UNCLEAN DATA:{report_content}")
            cleaned_data = self.clean_audit_data(report_content)
            
            # Process the document sections and insert evidence
            self._process_document(doc, cleaned_data, auditor_name)
            
            return doc
        except Exception as e:
            logger.error(f"Error filling template document: {str(e)}")
            raise
    
 

    def _fill_header_table(self, table, header_content):
        """Fill the header table with content."""
        if not table:
            return
            
        for row in table.rows:
            if len(row.cells) >= 2:
                field_name = row.cells[0].text.strip()
                if field_name in header_content:
                    row.cells[1].text = header_content[field_name]


    def _process_document(self, doc, content_sections, auditor_name):
        """Process all document sections and insert evidence in a single consolidated approach."""
        # Setup temp directory for images if needed
        if self.evidence_images and not self.report_images_dir:
            self.report_images_dir = os.path.join(tempfile.gettempdir(), f"report_images_{int(time.time())}")
            os.makedirs(self.report_images_dir, exist_ok=True)
        
        # logger.info(f"All image:{self.evidence_images} DATA:{content_sections}")
        # Extract the tables
        tables = doc.tables
        if len(tables) < 4:
            logger.error("Document doesn't have the expected table structure")
            return
            
        header_table = tables[0]
        legend_table = tables[1]
        evidence_table = tables[2]
        findings_table = tables[3]
        
        # Extract content from the clean structure
        header_content, legend_content, body_content, footer_content = {}, {}, [], {}
        nc_comments, ofi_comments = [], []
        
        # Process the content based on the clean structure
        if isinstance(content_sections, list) and len(content_sections) >= 3:
            if 'header' in content_sections[0]:
                header_content = content_sections[0]['header']
            if 'legend' in content_sections[0]:
                legend_content = content_sections[0]['legend']
            if 'body' in content_sections[1]:
                body_content = content_sections[1]['body']
            if 'footer' in content_sections[2]:
                footer_content = content_sections[2]['footer']
                # Remove specific keys
                for key in ['OK', 'NC', 'OFI', 'NA']:
                    if key in footer_content:
                        footer_content.pop(key)
        
        # 1. Fill the header table
        self._fill_header_table(header_table, header_content)
        
        # 2. Preprocess body_content to handle Excel files with multiple entries
        consolidated_content = self._consolidate_excel_entries(body_content)
        
        # 3. Fill the evidence table with consolidated entries
        nc_ofi_comments = self._fill_evidence_table(evidence_table, consolidated_content)
        
        # 4. Update footer content with collected NC/OFI comments
        if nc_ofi_comments['NC'] and 'NONCONFORMANCES' in footer_content:
            if footer_content['NONCONFORMANCES'] == 'Nil':
                footer_content['NONCONFORMANCES'] = "\n".join(nc_ofi_comments['NC'])
            else:
                footer_content['NONCONFORMANCES'] += "\n" + "\n".join(nc_ofi_comments['NC'])
                
        if nc_ofi_comments['OFI'] and 'OPPORTUNITIES FOR IMPROVEMENTS' in footer_content:
            if footer_content['OPPORTUNITIES FOR IMPROVEMENTS'] == 'Nil':
                footer_content['OPPORTUNITIES FOR IMPROVEMENTS'] = "\n".join(nc_ofi_comments['OFI'])
            else:
                footer_content['OPPORTUNITIES FOR IMPROVEMENTS'] += "\n" + "\n".join(nc_ofi_comments['OFI'])
        
        # 5. Fill the findings table
        self._fill_findings_table(findings_table, footer_content, auditor_name)

    def _consolidate_excel_entries(self, body_content):
        """Consolidate multiple entries with the same evidence file (especially Excel files)."""
        evidence_groups = {}
        consolidated_entries = []
        
        # Group entries by evidence file
        for entry in body_content:
            evidence_file = entry.get('SIGHTED EVIDENCE', '')
            process_name = entry.get('PROCESS', '')
            
            # Skip entries without evidence file
            if not evidence_file:
                consolidated_entries.append(entry)
                continue
                
            # Check if this is an Excel file
            is_excel = evidence_file.lower().endswith(('.xlsx', '.xls'))
            
            if is_excel:
                # For Excel files, group by evidence file
                if evidence_file not in evidence_groups:
                    evidence_groups[evidence_file] = {
                        'entries': [],
                        'base_entry': entry.copy(),  # Keep one entry as the base
                        'processes': []
                    }
                
                # Add to the group
                evidence_groups[evidence_file]['entries'].append(entry)
                evidence_groups[evidence_file]['processes'].append(process_name)
                
                # Take highest priority status (NC > OFI > OK > NA)
                for status in ['NC', 'OFI', 'OK', 'NA']:
                    if status in entry and entry[status] and not evidence_groups[evidence_file]['base_entry'].get(status):
                        evidence_groups[evidence_file]['base_entry'][status] = entry[status]
                
                # Combine comments if present
                if 'ADDITIONAL COMMENTS' in entry and entry['ADDITIONAL COMMENTS']:
                    if 'ADDITIONAL COMMENTS' not in evidence_groups[evidence_file]['base_entry'] or not evidence_groups[evidence_file]['base_entry']['ADDITIONAL COMMENTS']:
                        evidence_groups[evidence_file]['base_entry']['ADDITIONAL COMMENTS'] = entry['ADDITIONAL COMMENTS']
                    else:
                        evidence_groups[evidence_file]['base_entry']['ADDITIONAL COMMENTS'] += "; " + entry['ADDITIONAL COMMENTS']
            else:
                # Non-Excel files are added directly
                consolidated_entries.append(entry)
        
        # Add the consolidated Excel entries
        for evidence_file, group in evidence_groups.items():
            base_entry = group['base_entry']
            # Store all process names in a special field for use in _fill_row_with_evidence
            base_entry['_CONSOLIDATED_PROCESSES'] = group['processes']
            consolidated_entries.append(base_entry)
        
        return consolidated_entries

    def _fill_evidence_table(self, table, body_content):
        """Fill the evidence table with content and process images in one pass.
        Returns collected NC/OFI comments to be added to the footer."""
        if not table or len(table.rows) <= 1:
            return {'NC': [], 'OFI': []}
            
        # Clear all rows first
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            for cell_idx in range(len(row.cells)):
                row.cells[cell_idx].text = ""
        
        # Extract image text content once (avoid repeated calls)
        image_text_data = self._extract_image_text_content() if self.evidence_images else {}
        scores = self._extract_scores_from_text_data(image_text_data)

        
        # Track processes we've already added to avoid duplicates
        processed_entries = set()
        
        # Collect NC/OFI comments to be added to the footer
        nc_ofi_comments = {'NC': [], 'OFI': []}
        
        # Process each row
        for idx, entry in enumerate(body_content):
            if idx + 1 >= len(table.rows):  # Check bounds
                break
                    
            row = table.rows[idx + 1]
            if len(row.cells) < 7:  # Ensure we have all needed cells
                continue
                    
            # Get process and evidence data
            process_name = entry.get('PROCESS', '')
            evidence_file = entry.get('SIGHTED EVIDENCE', '')
            
            # Skip if we've already processed this process/evidence combination
            entry_key = f"{process_name}:{evidence_file}"
            if entry_key in processed_entries:
                logger.info(f"Skipping duplicate entry: {entry_key}")
                continue
                    
            processed_entries.add(entry_key)
            
            logger.info(f"Evidence passed: {evidence_file}")
            # Fill the row cells with content
            comments = self._fill_row_with_evidence(
                row, 
                process_name, 
                evidence_file, 
                image_text_data,
                entry,
                scores
            )
            
            # Collect NC/OFI comments for the footer
            if comments:
                # Record NC comments
                if 'NC' in comments and comments['NC']:
                    process_label = process_name
                    if '_CONSOLIDATED_PROCESSES' in entry:
                        process_label = evidence_file.split('.')[0]
                    nc_ofi_comments['NC'].append(f"{process_label}: {comments['NC']}")
                    
                # Record OFI comments
                if 'OFI' in comments and comments['OFI']:
                    process_label = process_name
                    if '_CONSOLIDATED_PROCESSES' in entry:
                        process_label = evidence_file.split('.')[0]
                    nc_ofi_comments['OFI'].append(f"{process_label}: {comments['OFI']}")
        
        return nc_ofi_comments

    def _fill_row_with_evidence(self, row, process_name, evidence_file, image_text_data, entry, scores):
        """Fill a row with evidence data and add image, checkmarks, and comments.
        Returns any NC/OFI comments for the footer."""
        logger.info(f"All files:{evidence_file}")
        if not process_name and not evidence_file:
            return None
        
        comments = {'NC': None, 'OFI': None}
            
        # Create process cell content
        if '_CONSOLIDATED_PROCESSES' in entry and entry['_CONSOLIDATED_PROCESSES'] and len(entry['_CONSOLIDATED_PROCESSES']) > 1:
            # For consolidated Excel entries, just show the file name as the process
            process_cell_content = f"Process: {evidence_file.split('.')[0]}"
        else:
            process_cell_content = f"Process: {process_name or evidence_file.split('.')[0]}"
        
        # For Excel files with multiple rows, consolidate all relevant data
        excel_data = []
        
        # Add extracted text data if available
        if evidence_file.lower().endswith(('.xlsx', '.xls')):
            # For Excel files, collect all customer data entries
            for key, val in image_text_data.items():
                # Look for matching evidence file
                if evidence_file.lower() in key.lower() or key.split(".")[0].lower() in evidence_file.lower():
                    excel_data.append(val)
            
            # Add consolidated Excel data if we found any
            if excel_data:
                process_cell_content += f"\n\nEvidence Data:\n" + "\n\n".join(excel_data)
                logger.info(f"Added Excel data for {evidence_file}")
        else:
            # For non-Excel files, just add the matching data
            for key, val in image_text_data.items():
                process_extract = key.split(".")[0]
                evidence = evidence_file.split(".")[0]
                if process_extract == evidence or process_extract == process_name:
                    process_cell_content += f"\n\nEvidence Data:\n{val}"
                    logger.info(f"Added data for File:{evidence_file} Process:{process_name}")
        
        # Add to process list if not already there
        if '_CONSOLIDATED_PROCESSES' in entry and entry['_CONSOLIDATED_PROCESSES']:
            for proc in entry['_CONSOLIDATED_PROCESSES']:
                clean_process_name = ' '.join(word.capitalize() for word in proc.replace('_', ' ').replace('-', ' ').split())
                if clean_process_name not in self.process_list:
                    self.process_list.append(clean_process_name)
        else:
            clean_process_name = ' '.join(word.capitalize() for word in process_name.replace('_', ' ').replace('-', ' ').split())
            if clean_process_name not in self.process_list:
                self.process_list.append(clean_process_name)
        
        # Fill basic cell data
        row.cells[0].text = process_cell_content
        row.cells[1].text = evidence_file
        
        # Find the evidence file and add image
        if self.evidence_images:
            matched_file = self._match_evidence_to_process(process_name, evidence_file)
            if matched_file:
                self._add_image_to_cell(row.cells[1], matched_file)
                
                # Determine category based on scores
                category, comment = self._determine_category(matched_file, scores.get(matched_file))
                
                # Add checkmarks and comments
                self._add_checkmark_with_background(row, category, comment)
                
                # Record comment for footer if NC or OFI
                if category == 'NC' and comment:
                    comments['NC'] = comment
                elif category == 'OFI' and comment:
                    comments['OFI'] = comment
        
        # Fill in OK/OFI/NC/NA cells from entry data if not already done
        for i, col in enumerate(['OK', 'OFI', 'NC', 'NA']):
            if col in entry and entry[col] and not row.cells[2 + i].text.strip():
                row.cells[2 + i].text = entry[col]
                
                # If we have checkmarks in OFI or NC columns and no comments recorded yet,
                # check for additional comments
                if (col == 'NC' or col == 'OFI') and entry[col] == '✓' and not comments[col]:
                    if 'ADDITIONAL COMMENTS' in entry and entry['ADDITIONAL COMMENTS']:
                        comments[col] = entry['ADDITIONAL COMMENTS']
        
        # Only add comments to the row for NC or OFI
        if ('NC' in entry and entry['NC']) or ('OFI' in entry and entry['OFI']):
            if 'ADDITIONAL COMMENTS' in entry and entry['ADDITIONAL COMMENTS'] and not row.cells[6].text.strip():
                row.cells[6].text = entry['ADDITIONAL COMMENTS']
        
        return comments


    def _add_image_to_cell(self, cell, file_name):
        """Add an evidence image to a table cell."""
        # Clear any existing content
        cell.text = ""
        
        # Get image from the file
        if not self.evidence_images.get(file_name):
            cell.text = f"No image available for {file_name}"
            return
            
        img_data = self.evidence_images[file_name][0]
        
        try:
            # Create image file
            ext = img_data.get('format', 'png').lower()
            image_filename = f"evidence_{os.path.basename(file_name)}.{ext}"
            image_path = os.path.join(self.report_images_dir, image_filename)
            
            # Write image to file
            with open(image_path, 'wb') as img_file:
                img_file.write(img_data['data'])
            
            # Save path to prevent garbage collection
            self.image_paths.append(image_path)
            
            # Add image to cell with consistent size
            paragraph = cell.add_paragraph()
            run = paragraph.add_run()
            
            width_inches = 2.0
            height_inches = 1.5
            width_emu = int(width_inches * 914400)  # 1 inch = 914400 EMUs
            height_emu = int(height_inches * 914400)
            
            run.add_picture(image_path, width=width_emu, height=height_emu)
            
            # Add caption
            caption = cell.add_paragraph()
            caption_run = caption.add_run(f"Evidence: {os.path.basename(file_name)}")
            caption_run.font.size = Pt(8)
            
        except Exception as e:
            logger.error(f"Error adding image: {str(e)}")
            cell.text = f"Error loading image: {str(e)}"
    
    def _add_checkmark_with_background(self, row, category, comment):
        """Add a checkmark to the appropriate cell with background color."""
        # Map column indices
        col_indices = {
            'OK': 2,
            'OFI': 3,
            'NC': 4,
            'NA': 5,
            'COMMENTS': 6
        }
        
        # Define colors for different categories
        colors = {
            'OK': "92D050",  # Light green
            'OFI': "8DB3E2",  # Light blue
            'NC': "FF0000",   # Red
            'NA': "FFFFFF"    # White
        }
        
        # Don't clear cells if they already have content
        # Only add checkmark if the cell is empty
        idx = col_indices.get(category)
        if idx is not None and idx < len(row.cells) and not row.cells[idx].text.strip():
            target_cell = row.cells[idx]
            
            # Add checkmark
            target_cell.text = "✓"
            
            # Set background color
            color = colors.get(category, "FFFFFF")
            try:
                # Set cell background color
                tcPr = target_cell._element.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                
                # Remove existing shading if present
                for old_shading in tcPr.findall('.//w:shd', namespaces=tcPr.nsmap):
                    tcPr.remove(old_shading)
                
                # Add new shading
                tcPr.append(shading)
                
                # Center the checkmark
                for paragraph in target_cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                logger.warning(f"Failed to set background color: {str(e)}")
        
        # Add comment if needed and the comments cell is empty
        comments_idx = col_indices.get('COMMENTS')
        if comment and category in ['OFI', 'NC'] and comments_idx < len(row.cells) and not row.cells[comments_idx].text.strip():
            row.cells[comments_idx].text = comment
    
    def _fill_findings_table(self, table, footer_content, auditor_name):
        logger.info(f'COMMENTS: {footer_content}')

        """Fill the findings table with content."""
        if not table:
            return
    
        
        # Use exactly the same logic as the old method that works well
        for row in table.rows:
            if len(row.cells) >= 2:
                field_name = row.cells[0].text.strip()
                content = next((v for k, v in footer_content.items() if field_name.startswith(k)), None)

                if content is not None:
                    content = content.replace("<br>", "\n")
                    row.cells[1].text = content
                if 'AUDIT REPORT FINAL COMMENTS' in field_name:
                    existing_text = row.cells[1].text
                    if existing_text and existing_text.__len__() > 5:
                        row.cells[1].text = existing_text
                    else:
                        row.cells[1].text = ''
                        row.cells[1].text = (
                            f"{existing_text}\n\n"
                            f"{auditor_name if auditor_name else 'INTERNAL AUDITOR'}"
                            f"{' INTERNAL AUDITOR' if auditor_name else ''}\n"
                            f"{datetime.now().strftime('%d/%m/%Y')}"
                        )



    def _format_audit_comments(self, cell, content, auditor_name):
        """Format the audit report final comments section."""
        try:
            # Check if content already ends with a date
            date_pattern = r"\d{1,2}/\d{1,2}/\d{4}$"
            if re.search(date_pattern, content.strip()):
                # The content already has a proper date format at the end
                cell.text = content
                return
                
            # Extract and format the comment sections
            comment_parts = content.split('.')
            
            # Handle case where there might not be sentences ending in periods
            if len(comment_parts) <= 1:
                comment_text = content
                auditor_date = ""
            else:
                # Combine all but the last part with periods (restore original text)
                comment_text = ".".join(comment_parts[:-1]) + "."
                auditor_date = comment_parts[-1].strip()
            
            # Extract auditor name and date if available
            match = re.match(r"^(.*?)(\d{1,2}/\d{1,2}/\d{4})$", auditor_date.strip())
            
            if match:
                name = ' '.join(match.group(1).split())
                name = re.sub(r'(<br>\s*)+', '\n', name).replace("<br>", "")
                date = match.group(2)
            else:
                name = auditor_name if auditor_name else 'Internal Auditor'
                date = datetime.now().strftime('%d/%m/%Y')
            
            # Format the final content
            cell.text = f"{comment_text}\n\n{name}\n{date}"
        except Exception as e:
            logger.error(f"Error formatting audit comments: {str(e)}")
            # Fallback to simple format
            cell.text = f"{content}\n\n{auditor_name if auditor_name else 'Internal Auditor'}\n{datetime.now().strftime('%d/%m/%Y')}"
    
    def _extract_image_text_content(self):
        """Extract text content from evidence images using GPT-4o Vision."""
        if not self.evidence_images or not self.openai_api_key:
            return {}
            
        extracted_contents = {}
        
        try:
            import openai
            client = openai.OpenAI(api_key=self.openai_api_key)
            
            for evidence_file, images in self.evidence_images.items():
                if not images:
                    continue
                
                # Check if this is an Excel file - handle differently
                is_excel = evidence_file.lower().endswith(('.xlsx', '.xls'))
                
                # Specific prompt for different file types    
                system_prompt = """You are an expert at extracting customer feedback data. 
                Extract ONLY the following key fields in this EXACT format:
                
                **Customer:** [Customer Name]
                **Date:** [Date in DD/MM/YYYY format]
                **Score:** [Total Score, format as X/Y]  
                **Comments:** [Brief summary of key comments]
                
                If you cannot find one of these fields, leave it blank but maintain the format.
                Be concise and precise."""
                
                # Different prompt for Excel files
                if is_excel:
                    system_prompt = """You are an expert at extracting data from customer feedback spreadsheets. 
                    This image contains a spreadsheet with multiple customer feedback entries.
                    
                    Extract data for ALL VISIBLE customers in the image, with each customer's data in this format:
                    IN THE CASE WHERE CUSTOMER DATA IS NOT IN THE IMAGE BUT A PROCESS. EXTRACT THE PROCESS DON'T FAIL 
                    TO EXTRACT
                    
                    **Customer:** [Customer Name]
                    **Date:** [Date in DD/MM/YYYY format]
                    **Score:** [Total Score, format as X/Y]  
                    **Comments:** [Brief summary of key comments]

                    REMEMBER FOR DIAGRAM IMAGES ONLY. ALL DATA IS IN DIAGRAMS LOOK THROUGH THE DIAGRAM THE TITLE BECOMES CUSTOMER NAME
                    DATE IS ALWAYS AT THE FOOTER ELSE ADD NONE, IF NOTHING LIKE SCORE ADD IN ITALIC IMAGES CAN'T BE SCORED AND 
                    NOW GIVE ME BRIEF COMMENTS FOR THE IMAGE AS YOU'RE DOING ABOVE. EVEN IF IT'S NOT CUSTOMER DATA, IF ITS A PROCESS GIVE DETAILS OF IT EXTRACTING ALL INFO
                    IF IMAGE DOEST NOT CONTAIN DATA, THE PROCESS IS THE EVIDENCE DATA IN THE IMAGE RETURN IT
                    
                    Present each customer as a separate entry with a blank line between entries.
                    If data for any field is missing, leave it blank but maintain the format.
                    Be concise and precise."""
                    
                try:
                    # Use first image
                    img_data = images[0]
                    
                    # Convert to Base64
                    import base64
                    base64_image = base64.b64encode(img_data['data']).decode('utf-8')
                    
                    # Create prompt for GPT-4o Vision
                    messages = [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": f"Extract the [customer feedback data] / [data process] from this {'spreadsheet' if is_excel else 'image'}."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/{img_data.get('format', 'png')};base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ]
                    
                    # Call the API
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages,
                        max_tokens=1000  # Increased for Excel files with multiple entries
                    )
                    
                    # Extract the content
                    extracted_text = response.choices[0].message.content
                    extracted_contents[evidence_file] = extracted_text
                    logger.info(f"Successfully extracted text from {os.path.basename(evidence_file)}")
                    
                except Exception as e:
                    logger.error(f"Error extracting from {evidence_file}: {str(e)}")
                    
        except Exception as e:
            logger.error(f"Error setting up image extraction: {str(e)}")
            
        return extracted_contents
    
    def _extract_scores_from_text_data(self, image_text):
        """Extract scores from image text data."""
        scores = {}
        
        if not image_text:
            return scores
            
        for company, text in image_text.items():
            # Look for score pattern - handle both Score: X/Y and Score: X formats
            score_match = re.search(r'\*\*Score:\*\* (\d+)(?:/\d+)?', text)
            if score_match:
                try:
                    score_value = int(score_match.group(1))
                    scores[company] = score_value
                except (ValueError, IndexError):
                    pass
            
            # Also check for alternative score formats that might appear in Excel data
            alt_score_match = re.search(r'score:\s*(\d+(?:\.\d+)?)', text, re.IGNORECASE)
            if not score_match and alt_score_match:
                try:
                    score_value = float(alt_score_match.group(1))
                    scores[company] = score_value
                except (ValueError, IndexError):
                    pass
        
        return scores
    
    def _determine_category(self, file_name, score=None):
        """Determine category (OK, OFI, NC, NA) based on score information."""
        # Default to OK for Excel files
        if file_name.lower().endswith(('.xlsx', '.xls')):
            return 'OK', ""
        
        # If no score data, default to OK
        if not score:
            return 'OK', ""
        
        try:
            percentage = 0
            
            # Handle different score formats
            if isinstance(score, (int, float)):
                # Convert to percentage based on scale
                if score <= 10:  # 0-10 scale
                    percentage = score * 10
                elif score <= 25:  # Appears to be out of 25 based on example
                    percentage = (score / 25) * 100
                elif score <= 100:  # Already a percentage
                    percentage = score
                else:
                    # Invalid score
                    return 'OK', ""
            # Handle string scores in format X/Y
            elif isinstance(score, str) and '/' in score:
                parts = score.split('/')
                try:
                    numerator = float(parts[0].strip())
                    denominator = float(parts[1].strip())
                    if denominator > 0:
                        percentage = (numerator / denominator) * 100
                except (ValueError, IndexError):
                    # Invalid score format
                    return 'OK', ""
            else:
                # Unknown score format
                return 'OK', ""
                
            # Determine category based on percentage
            if percentage >= 90:
                return 'OK', ""
            elif percentage >= 80:
                return 'OK', ""
            elif percentage >= 70:
                score_display = f"{score}/25" if isinstance(score, (int, float)) and score <= 25 else score
                return 'OFI', f"Score of {score_display} indicates room for improvement in customer satisfaction. Consider implementing additional feedback mechanisms."
            else:
                score_display = f"{score}/25" if isinstance(score, (int, float)) and score <= 25 else score
                return 'NC', f"Low score of {score_display} represents a significant gap in customer satisfaction requiring immediate corrective action."
                
        except Exception as e:
            logger.error(f"Error determining category: {str(e)}")
            return 'OK', ""
    
    def save_report(self, doc, output_path, convert_to_pdf=True):
        """Save the document to the specified path and optionally convert to PDF."""
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # Determine file paths
            base_path = output_path.rsplit('.', 1)[0] if '.' in output_path else output_path
            docx_path = f"{base_path}.docx"
            pdf_path = f"{base_path}.pdf" if convert_to_pdf else None
            
            # Save the DOCX file
            doc.save(docx_path)
            logger.info(f"Report saved to {docx_path}")
            
            # Convert to PDF if requested
            if convert_to_pdf:
                pdf_success = self._convert_to_pdf(docx_path, pdf_path)
                
                if pdf_success:
                    logger.info(f"PDF created at {pdf_path}")
                    return True
                else:
                    logger.error("Failed to create PDF")
                    return False
            
            return True
        
        except Exception as e:
            logger.error(f"Error saving report: {str(e)}")
            return False
        
    def _match_evidence_to_process(self, process_name, evidence_file):
        """Find the most relevant evidence file for a process."""
        if not self.evidence_images:
            return None
            
        evidence_files = list(self.evidence_images.keys())
        if not evidence_files:
            return None
            
        # First try: Match based on exact evidence file name
        if evidence_file:
            for file in evidence_files:
                if file in self.used_evidence:
                    continue
                
                # Check if the evidence file name is contained in the full file path
                if evidence_file.lower() in os.path.basename(file).lower():
                    self.used_evidence.add(file)
                    return file
        
        # Extract keywords from process text
        process_keywords = set()
        if process_name:
            process_keywords = set(process_name.lower().replace('.', ' ').replace('_', ' ').replace('-', ' ').split())
        
        evidence_keywords = set()
        if evidence_file:
            evidence_keywords = set(evidence_file.lower().replace('.', ' ').replace('_', ' ').replace('-', ' ').split())
        
        # Second try: Look for keyword overlap with both process and evidence name
        for file in evidence_files:
            if file in self.used_evidence:
                continue
                
            file_base = os.path.basename(file).lower()
            file_keywords = set(file_base.replace('.', ' ').replace('_', ' ').replace('-', ' ').split())
            
            # Check for keyword overlap with process or evidence name
            if (process_keywords and process_keywords.intersection(file_keywords)) or (evidence_keywords and evidence_keywords.intersection(file_keywords)):
                self.used_evidence.add(file)
                return file
        
        # Third try: Look for any keyword match
        for file in evidence_files:
            if file in self.used_evidence:
                continue
                
            file_base = os.path.basename(file).lower()
            
            for keyword in process_keywords.union(evidence_keywords):
                if keyword in file_base and len(keyword) > 2:  # Only match on meaningful keywords
                    self.used_evidence.add(file)
                    return file
        
        # Final try: Just use any unused file
        for file in evidence_files:
            if file not in self.used_evidence:
                self.used_evidence.add(file)
                return file
        
        # If all files are used, return the first one that best matches
        # Score each file for relevance
        best_score = -1
        best_file = None
        
        for file in evidence_files:
            file_base = os.path.basename(file).lower()
            file_keywords = set(file_base.replace('.', ' ').replace('_', ' ').replace('-', ' ').split())
            
            # Calculate score based on keyword overlap
            score = len(process_keywords.intersection(file_keywords)) + len(evidence_keywords.intersection(file_keywords))
            
            # Bonus for exact matches
            if process_name and process_name.lower() in file_base:
                score += 10
            if evidence_file and evidence_file.lower() in file_base:
                score += 10
                
            if score > best_score:
                best_score = score
                best_file = file
        
        # If we found a good match, use it even if used before
        if best_file and best_score > 0:
            return best_file
            
        # Otherwise, just return the first file
        return evidence_files[0] if evidence_files else None

    def _convert_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF using available methods."""
        # Method 1: Try with win32com (Windows only)
        try:
            import win32com.client
            pythoncom.CoInitialize()
            
            # Initialize Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # Open document and save as PDF
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF
                doc.Close()
                word.Quit()
                
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    return True
            finally:
                try:
                    word.Quit()
                except:
                    pass
                pythoncom.CoUninitialize()
                
        except Exception as e:
            logger.warning(f"Error using win32com for PDF conversion: {str(e)}")
        
        # Method 2: Try with docx2pdf if available
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return True
        except ImportError:
            logger.info("docx2pdf not available")
        except Exception as e:
            logger.warning(f"Error using docx2pdf: {str(e)}")
        
        # No methods succeeded
        return False